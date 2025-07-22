from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import xml.etree.ElementTree as ET

class DocumentStandardsParser:
    """คลาสสำหรับอ่าน XML standards"""
    
    @staticmethod
    def parse_xml_file(xml_file_path):
        """อ่านไฟล์ XML และแปลงเป็น dictionary"""
        try:
            tree = ET.parse(xml_file_path)
            root = tree.getroot()
            return DocumentStandardsParser._parse_xml_root(root)
        except Exception as e:
            print(f"Error reading XML file: {e}")
            return DocumentStandardsParser._get_default_standards()
    
    @staticmethod
    def parse_xml_content(xml_content):
        """อ่าน XML content และแปลงเป็น dictionary"""
        try:
            root = ET.fromstring(xml_content)
            return DocumentStandardsParser._parse_xml_root(root)
        except Exception as e:
            print(f"Error parsing XML content: {e}")
            return DocumentStandardsParser._get_default_standards()
    
    @staticmethod
    def _parse_xml_root(root):
        """แปลง XML root element เป็น dictionary"""
        return {
            'theme_font': root.find('.//theme_font').text,
            'main_heading_size': int(root.find('.//main_heading').text),
            'content_size': int(root.find('.//content').text),
            'sub_heading_size': int(root.find('.//sub_heading').text),
            'cover_title_size': int(root.find('.//cover_title').text),
            'cover_info_size': int(root.find('.//cover_info').text),
            
            # Formatting rules
            'main_heading_bold': root.find('.//main_heading_bold').text.lower() == 'true',
            'sub_heading_bold': root.find('.//sub_heading_bold').text.lower() == 'true',
            'content_bold': root.find('.//content_bold').text.lower() == 'true',
            'cover_title_bold': root.find('.//cover_title_bold').text.lower() == 'true',
            
            # เพิ่มรองรับ major_heading (หัวข้อใหญ่) ที่ใช้ขนาด sub_heading
            'major_heading_size': int(root.find('.//sub_heading').text),
            'major_heading_bold': root.find('.//sub_heading_bold').text.lower() == 'true',
            
            # Indentation rules
            'paragraph_indent': root.find('.//paragraph_indent').text.lower() == 'true',
            'sub_heading_indent': root.find('.//sub_heading_indent').text.lower() == 'true',
            
            # Alignment rules
            'cover_center': root.find('.//cover_center').text.lower() == 'true',
            'main_heading_left': root.find('.//main_heading_left').text.lower() == 'true',
            'content_justify': root.find('.//content_justify').text.lower() == 'true',
            
            # References
            'picture_position': root.find('.//picture_position').text if root.find('.//picture_position') is not None else 'below_center',
            'table_position': root.find('.//table_position').text if root.find('.//table_position') is not None else 'above_left',
            
            # Section lists
            'main_headings': [item.text for item in root.findall('.//main_headings/item')],
            'cover_sections': [item.text for item in root.findall('.//cover_sections/item')],
            'no_indent_sections': [item.text for item in root.findall('.//no_indent_sections/item')],
            'center_sections': [item.text for item in root.findall('.//center_sections/item')],
            'sub_heading_patterns': [item.text for item in root.findall('.//sub_heading_patterns/pattern')],
            'special_spacing_sections': [item.text for item in root.findall('.//special_spacing_sections/item')] if root.find('.//special_spacing_sections') is not None else [],
            
            # Exceptions
            'skip_font_check': [item.text for item in root.findall('.//skip_font_check/item')],
            'skip_bold_check': [item.text for item in root.findall('.//skip_bold_check/item')]
        }
    
    @staticmethod
    def _get_default_standards():
        """ค่า default หากไม่สามารถอ่าน XML ได้"""
        return {
            'theme_font': 'TH Sarabun New',
            'main_heading_size': 18,  # ชื่อบท
            'major_heading_size': 16,  # หัวข้อใหญ่ (1. 2. 3.)
            'content_size': 14,
            'sub_heading_size': 16,   # เก็บไว้เพื่อความเข้ากันได้
            'cover_title_size': 18,
            'cover_info_size': 14,
            'main_heading_bold': True,
            'major_heading_bold': True,  # หัวข้อใหญ่เป็น Bold
            'sub_heading_bold': True,  # เก็บไว้เพื่อความเข้ากันได้
            'content_bold': False,
            'cover_title_bold': True,
            'paragraph_indent': True,
            'sub_heading_indent': False,
            'cover_center': True,
            'main_heading_left': True,
            'content_justify': False,
            'picture_position': 'below_center',
            'table_position': 'above_left',
            'main_headings': ['บทคัดย่อ', 'Abstract', 'กิตติกรรมประกาศ', 'สารบัญ', 'บทที่ 1', 'บทที่ 2', 'บทที่ 3', 'บทที่ 4', 'บทที่ 5', 'เอกสารอ้างอิง', 'ภาคผนวก', 'ประวัติผู้เขียน'],
            'cover_sections': ['เอกสารโครงงานฉบับสมบูรณ์', 'คู่มือการใช้งานระบบ', 'CS/IT', 'โดย', 'อาจารย์ที่ปรึกษา'],
            'no_indent_sections': ['ปก', 'สารบัญ', 'สารบัญตาราง', 'สารบัญภาพ', 'บทคัดย่อ', 'Abstract', 'กิตติกรรมประกาศ'],
            'center_sections': ['ปก', 'ภาคผนวก'],
            'sub_heading_patterns': [r'^\d+\.\d+\s', r'^\d+\.\d+\.\d+\s', r'^\(\d+\)\s'],
            'special_spacing_sections': ['บทคัดย่อ', 'Abstract'],
            'skip_font_check': ['หมายเลขหน้า', 'Header', 'Footer'],
            'skip_bold_check': ['Abstract', 'บทคัดย่อ', 'Acknowledgement']
        }

class DocumentChecker:
    """คลาสหลักสำหรับตรวจสอบเอกสาร"""
    
    def __init__(self, standards):
        self.standards = standards
        self.errors = []
        self.tables = []
        self.pictures = []
    
    def check_document(self, doc_path):
        """ตรวจสอบเอกสาร Word"""
        doc = Document(doc_path)
        self.errors = []
        self.tables = []
        self.pictures = []

        # ตรวจหาตารางและรูปภาพในเอกสาร
        self._extract_tables_and_pictures(doc)

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # ข้ามย่อหน้าว่าง
                # ข้ามหน้าปกในการตรวจสอบปกติ
                if not self._is_cover_page_section(paragraph.text, i):
                    self._check_paragraph(paragraph, i + 1, i)

        # ตรวจสอบการอ้างอิงตารางและรูปภาพ
        self._check_references(doc)

        return self.errors, doc

    def _extract_tables_and_pictures(self, doc):
        """ดึงข้อมูลตารางและรูปภาพจากเอกสาร"""
        table_count = 0
        for table in doc.tables:
            table_count += 1
            self.tables.append({
                'number': table_count,
                'position': None
            })

    def _is_main_heading(self, text, paragraph=None):
        """ตรวจสอบว่าเป็นชื่อบท (18pt Bold) หรือไม่"""
        text_clean = text.strip()
        
        # 1. ตรวจสอบหัวข้อบท - บทที่ X
        if re.match(r'^บทที่\s*\d+', text_clean):
            return True
        
        # 2. ตรวจสอบชื่อบทที่รู้จัก (ไม่ขึ้นต้นด้วยตัวเลข)
        chapter_names = [
            'บทนำ', 'งานวิจัยและทฤษฎีที่เกี่ยวข้อง', 'การวิเคราะห์และออกแบบระบบ',
            'การพัฒนาระบบ', 'การทดสอบระบบ', 'สรุปและข้อเสนอแนะ',
            'เป้าหมายและขอบเขต', 'ปัญหาและขอบเขต'
        ]
        
        # ตรวจสอบว่าไม่ขึ้นต้นด้วยตัวเลข และมีชื่อบทที่รู้จัก
        if not re.match(r'^\d+\.', text_clean):
            for chapter_name in chapter_names:
                if chapter_name in text_clean:
                    return True
        
        # 3. ตรวจสอบหัวข้อหลักอื่นๆ ที่กำหนดใน standards (ไม่ขึ้นต้นด้วยตัวเลข)
        main_headings_extended = self.standards['main_headings'] + [
            'สารบัญตาราง', 'สารบัญภาพ', 'รายการสัญลักษณ์และคำย่อ',
            'บทที่ 1', 'บทที่ 2', 'บทที่ 3', 'บทที่ 4', 'บทที่ 5',
            'References', 'Bibliography', 'Appendix'
        ]
        
        # ตรวจสอบว่าไม่ขึ้นต้นด้วยตัวเลข และมีหัวข้อที่กำหนด
        if not re.match(r'^\d+\.', text_clean):
            for heading in main_headings_extended:
                if heading in text_clean:
                    return True
        
        # 4. ตรวจสอบข้อความที่อาจเป็นชื่อบท (สำหรับชื่อบทที่จัดแนวกึ่งกลาง)
        if (paragraph and 
            paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and 
            len(text_clean) < 50 and  # ชื่อบทมักจะไม่ยาวมาก
            not re.match(r'^\d+\.', text_clean) and  # ไม่ขึ้นต้นด้วยเลข
            not self._is_cover_page_section(text_clean, 0)):  # ไม่ใช่หน้าปก
            return True
        
        return False

    def _is_major_heading(self, text, paragraph=None):
        """ตรวจสอบว่าเป็นหัวข้อใหญ่หรือไม่ - รูปแบบ 1. 2. 3. ขนาด 16pt Bold"""
        text_clean = text.strip()
        
        # ตรวจสอบรูปแบบหัวข้อใหญ่ 1. 2. 3. (ไม่ใช่ 1.1 หรือ 1.1.1)
        if re.match(r'^\d+\.\s+[^\d]', text_clean):  # เริ่มด้วยตัวเลข. ตามด้วยช่องว่างและไม่ใช่ตัวเลข
            return True
        
        return False

    def _is_sub_heading_level1(self, text):
        """ตรวจสอบว่าเป็นหัวข้อย่อยระดับ 1 หรือไม่ - รูปแบบ 1.1 1.2 2.1 ขนาด 14pt ปกติ"""
        text_clean = text.strip()
        
        # รูปแบบ 1.1 1.2 2.1 (มีจุดสองจุด)
        if re.match(r'^\d+\.\d+\s', text_clean):
            return True
        
        return False

    def _is_sub_heading_level2(self, text):
        """ตรวจสอบว่าเป็นหัวข้อย่อยระดับ 2 หรือไม่ - รูปแบบ 1.1.1 2.2.1 ขนาด 14pt ปกติ"""
        text_clean = text.strip()
        
        # รูปแบบ 1.1.1 2.2.1 (มีจุดสามจุด)
        if re.match(r'^\d+\.\d+\.\d+\s', text_clean):
            return True
        
        return False

    def _is_sub_heading_level3(self, text):
        """ตรวจสอบว่าเป็นหัวข้อย่อยระดับ 3 หรือไม่ - รูปแบบ (1) (2) (3) ขนาด 14pt ปกติ"""
        text_clean = text.strip()
        
        # รูปแบบ (1) (2) (3)
        if re.match(r'^\(\d+\)\s', text_clean):
            return True
        
        return False

    def _is_any_sub_heading(self, text):
        """ตรวจสอบว่าเป็นหัวข้อย่อยระดับใดๆ หรือไม่"""
        return (self._is_sub_heading_level1(text) or 
                self._is_sub_heading_level2(text) or 
                self._is_sub_heading_level3(text))

    def _get_sub_heading_level(self, text):
        """ดึงระดับของหัวข้อย่อย (1, 2, หรือ 3)"""
        if self._is_sub_heading_level1(text):
            return 1
        elif self._is_sub_heading_level2(text):
            return 2
        elif self._is_sub_heading_level3(text):
            return 3
        return 0

    def _check_thai_font(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบฟอนต์ภาษาไทย"""
        text = paragraph.text.strip()
        if not text or self._should_skip_font_check(text):
            return
        
        # ข้ามถ้ามี error message อยู่แล้ว
        if "[❌ ข้อผิดพลาด:" in text:
            return
        
        expected_font = self.standards['theme_font']
        font_errors = []
        
        # ตรวจสอบทุก run
        for i, run in enumerate(paragraph.runs):
            if not run.text.strip():
                continue
            
            # ตรวจสอบฟอนต์
            if run.font.name and not self._font_matches(run.font.name, expected_font):
                font_errors.append(run.font.name)
                run.font.highlight_color = 7  # Yellow highlight
        
        # ถ้าพบข้อผิดพลาด
        if font_errors:
            detected_fonts = list(set(font_errors))
            fonts_str = ', '.join(detected_fonts)
            error_message = f'ฟอนต์ไม่ถูกต้อง: ควรใช้ "{expected_font}" แต่พบ "{fonts_str}"'
            
            self.errors.append({
                'paragraph': para_num,
                'type': 'ฟอนต์ไม่ถูกต้อง',
                'description': error_message,
                'paragraph_obj': paragraph
            })
            
            self._highlight_paragraph(paragraph)
            self._add_comment_to_paragraph(paragraph, error_message)

    def _font_matches(self, actual_font, expected_font):
        """เปรียบเทียบฟอนต์ว่าตรงกันหรือไม่"""
        if not actual_font:
            return False
        
        actual_lower = str(actual_font).lower().strip()
        expected_lower = str(expected_font).lower().strip()
        
        # ตรวจสอบการตรงกันแบบตรงไปตรงมา
        if actual_lower == expected_lower:
            return True
        
        # ตรวจสอบ variant ของ TH Sarabun New
        if 'sarabun' in expected_lower and 'sarabun' in actual_lower:
            return True
        
        return False

    def _check_indentation_detailed(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบการเยื้องแบบละเอียด - รองรับหัวข้อย่อยหลายระดับ"""
        text = paragraph.text.strip()
        if not text:
            return
        
        # ข้ามการตรวจสอบสำหรับ caption
        if self._is_caption_style(paragraph):
            return
        
        # ข้ามการตรวจสอบสำหรับส่วนที่ไม่ต้องเยื้อง
        if self._should_skip_indent_check(text):
            return
        
        # ข้ามส่วน Keywords ใน Abstract
        if self._is_keywords_section(text):
            return
        
        # ข้ามเลขหน้าและคำว่า "ผู้จัดทำ" ในสารบัญ
        if self._is_table_of_contents_page_number(text):
            return
        
        # ข้ามชื่อบทและหัวข้อใหญ่
        if self._is_main_heading(text, paragraph) or self._is_major_heading(text, paragraph):
            return
        
        # ข้ามการอ้างอิงตารางและรูปภาพ (ในเนื้อหา ไม่ใช่ caption)
        if self._contains_table_reference(text) or self._contains_picture_reference(text):
            return
        
        # ข้ามเนื้อหาในกิตติกรรมประกาศ
        if self._is_acknowledgement_content(text):
            return
        
        # ตรวจสอบการเยื้องสำหรับหัวข้อย่อยแต่ละระดับ
        sub_heading_level = self._get_sub_heading_level(text)
        
        if sub_heading_level > 0:
            # หัวข้อย่อย - ตรวจสอบการเยื้องตามระดับ
            if sub_heading_level == 1:
                expected_indent_inches = 0.0
            elif sub_heading_level == 2:
                expected_indent_inches = 0.5
            elif sub_heading_level == 3:
                expected_indent_inches = 1.5
            else:
                expected_indent_inches = sub_heading_level * 0.5
            
            left_indent = paragraph.paragraph_format.left_indent
            actual_indent_inches = 0
            
            if left_indent and hasattr(left_indent, 'inches'):
                actual_indent_inches = left_indent.inches
            
            # ให้ความผิดพลาด ±0.2 นิ้ว
            tolerance = 0.2
            if abs(actual_indent_inches - expected_indent_inches) > tolerance:
                level_description = {
                    1: "ระดับ 1 (1.1, 2.1)",
                    2: "ระดับ 2 (1.1.1, 2.1.1)", 
                    3: "ระดับ 3 ((1), (2))"
                }
                
                desc = level_description.get(sub_heading_level, f"ระดับ {sub_heading_level}")
                
                if sub_heading_level == 3:
                    error_msg = f'หัวข้อย่อย{desc} ควรกด Tab 3 ครั้ง (เยื้อง {expected_indent_inches} นิ้ว) แต่พบ {actual_indent_inches:.2f} นิ้ว'
                    comment_msg = f"หัวข้อย่อย{desc} ต้องกด Tab 3 ครั้ง (ปัจจุบัน {actual_indent_inches:.2f} นิ้ว)"
                else:
                    error_msg = f'หัวข้อย่อย{desc} ควรเยื้อง {expected_indent_inches} นิ้ว (พบ {actual_indent_inches:.2f} นิ้ว)'
                    comment_msg = f"หัวข้อย่อย{desc} ควรเยื้อง {expected_indent_inches} นิ้ว (ปัจจุบัน {actual_indent_inches:.2f} นิ้ว)"
                
                self.errors.append({
                    'paragraph': para_num,
                    'type': 'การเยื้องหัวข้อย่อยไม่ถูกต้อง',
                    'description': error_msg,
                    'paragraph_obj': paragraph
                })
                self._highlight_paragraph(paragraph)
                self._add_comment_to_paragraph(paragraph, comment_msg)
        
        else:
            # เนื้อหาทั่วไป - ควรมีการเยื้องย่อหน้า (first line indent)
            if self.standards['paragraph_indent']:
                first_line_indent = paragraph.paragraph_format.first_line_indent
                
                should_indent = (first_line_indent is None or (first_line_indent and hasattr(first_line_indent, 'inches') and first_line_indent.inches < 0.3))
                
                if should_indent:
                    self.errors.append({
                        'paragraph': para_num,
                        'type': 'การเยื้องย่อหน้าไม่ถูกต้อง',
                        'description': 'เนื้อหาควรมีการเยื้องย่อหน้า (First Line Indent ประมาณ 0.5 นิ้ว)',
                        'paragraph_obj': paragraph
                    })
                    self._highlight_paragraph(paragraph)
                    self._add_comment_to_paragraph(paragraph, "เนื้อหาควรมีการเยื้องย่อหน้า (กด Tab หรือตั้งค่า First Line Indent 0.5 นิ้ว)")

    # ฟังก์ชันช่วยเหลือ
    def _is_acknowledgement_content(self, text):
        """ตรวจสอบว่าเป็นเนื้อหาในกิตติกรรมประกาศหรือไม่"""
        acknowledgement_content_patterns = [
            'ขอขอบพระคุณ', 'ขอขอบคุณ', 'ขอกราบขอบพระคุณ', 
            'คุณพ่อ', 'คุณแม่', 'ที่ให้คำปรึกษา', 'ที่ให้การสนับสนุน',
            'ครอบครัว', 'เพื่อน', 'อาจารย์', 'ที่ปรึกษา', 'คณาจารย์',
            'แนวทางในการ', 'ให้คำแนะนำ', 'และการพัฒนา', 'ในการเรียน',
            'ผู้เขียน', 'คณะผู้จัดทำ', 'ขอกราบ'
        ]
        
        for pattern in acknowledgement_content_patterns:
            if pattern in text:
                return True
        return False

    def _is_keywords_section(self, text):
        """ตรวจสอบว่าเป็นส่วน Keywords หรือไม่"""
        keywords_patterns = [
            'keywords:', 'keyword:', 'คำสำคัญ:', 'คำสำคัญ :', 'key words:'
        ]
        text_lower = text.lower().strip()
        
        for pattern in keywords_patterns:
            if text_lower.startswith(pattern):
                return True
        return False

    def _is_table_of_contents_page_number(self, text):
        """ตรวจสอบว่าเป็นเลขหน้าในสารบัญหรือไม่"""
        text_clean = text.strip()
        
        # เลขหน้าที่อยู่ท้ายบรรทัดในสารบัญ
        if re.match(r'^\d+', text_clean):
            return True
        
        # คำว่า "หน้า" หรือ "ผู้จัดทำ" ที่อยู่ในสารบัญ
        if text_clean in ['หน้า', 'ผู้จัดทำ', 'Page']:
            return True
            
        return False

    def _check_alignment(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบการจัดแนว"""
        text = paragraph.text.strip()
        if not text:
            return
        
        # ข้ามการตรวจสอบสำหรับ caption
        if self._is_caption_style(paragraph):
            return
        
        # ตรวจสอบการจัดแนวสำหรับชื่อบท
        if self._is_main_heading(text, paragraph):
            # ชื่อบทควรจัดแนวกึ่งกลาง หรือ ชิดซ้าย (ยืดหยุ่นตามรูปแบบ)
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                self.errors.append({
                    'paragraph': para_num,
                    'type': 'การจัดแนวไม่ถูกต้อง',
                    'description': 'ชื่อบทควรจัดแนวกึ่งกลางหรือชิดซ้าย (ไม่ใช่ชิดขวา)',
                    'paragraph_obj': paragraph
                })
                self._highlight_paragraph(paragraph)
                self._add_comment_to_paragraph(paragraph, "ชื่อบทควรจัดแนวกึ่งกลางหรือชิดซ้าย")
        
        # หัวข้อใหญ่ และหัวข้อย่อยควรชิดซ้าย
        elif ((self._is_major_heading(text, paragraph) or self._is_any_sub_heading(text)) 
            and self.standards['main_heading_left']):
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER or paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                self.errors.append({
                    'paragraph': para_num,
                    'type': 'การจัดแนวไม่ถูกต้อง',
                    'description': 'หัวข้อควรจัดแนวชิดซ้าย',
                    'paragraph_obj': paragraph
                })
                self._highlight_paragraph(paragraph)
                self._add_comment_to_paragraph(paragraph, "หัวข้อควรจัดแนวชิดซ้าย")

    def _check_font_sizes(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบขนาดฟอนต์"""
        text = paragraph.text.strip()
        if not text:
            return
        
        # ข้ามการตรวจสอบสำหรับ caption
        if self._is_caption_style(paragraph):
            return
            
        expected_size = None
        section_type = None
        
        # 1. ชื่อบท - 18pt Bold
        if self._is_main_heading(text, paragraph):
            expected_size = 18
            section_type = "ชื่อบท"
        # 2. หัวข้อใหญ่ (1. 2. 3.) - 16pt Bold
        elif self._is_major_heading(text, paragraph):
            expected_size = 16
            section_type = "หัวข้อใหญ่"
        # 3. หัวข้อย่อยทุกระดับ (1.1, 1.1.1, (1)) - 14pt ปกติ
        elif self._is_any_sub_heading(text):
            expected_size = 14
            section_type = "หัวข้อย่อย"
        # 4. เนื้อหาทั่วไป - 14pt ปกติ
        else:
            expected_size = 14
            section_type = "เนื้อหา"
        
        # ตรวจสอบขนาดฟอนต์ในทุก run
        has_size_error = False
        wrong_sizes = []
        
        for run in paragraph.runs:
            run_text = run.text.strip()
            if not run_text or "[❌ ข้อผิดพลาด:" in run_text:
                continue
                
            actual_size = None
            
            try:
                if run.font.size and hasattr(run.font.size, 'pt'):
                    actual_size = run.font.size.pt
                elif paragraph.style and paragraph.style.font and paragraph.style.font.size:
                    if hasattr(paragraph.style.font.size, 'pt'):
                        actual_size = paragraph.style.font.size.pt
            except AttributeError:
                continue
            
            # ให้ความผิดพลาด ±1pt สำหรับความยืดหยุ่น
            if actual_size and abs(actual_size - expected_size) > 1:
                has_size_error = True
                if actual_size not in wrong_sizes:
                    wrong_sizes.append(actual_size)
        
        # ถ้าพบข้อผิดพลาด ให้เพิ่มข้อความแจ้งเตือน
        if has_size_error and wrong_sizes:
            size_list = ", ".join([f"{size}pt" for size in wrong_sizes])
            self.errors.append({
                'paragraph': para_num,
                'type': 'ขนาดฟอนต์ไม่ถูกต้อง',
                'description': f'{section_type}ควรใช้ขนาด {expected_size}pt แต่พบขนาด {size_list}',
                'paragraph_obj': paragraph
            })
            self._highlight_paragraph(paragraph)
            self._add_comment_to_paragraph(paragraph, f"{section_type}ควรใช้ขนาด {expected_size}pt (พบขนาด {size_list})")

    def _check_font_bold(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบความหนาของฟอนต์"""
        text = paragraph.text.strip()
        if not text or self._should_skip_bold_check(text):
            return
        
        # ข้ามการตรวจสอบสำหรับ caption
        if self._is_caption_style(paragraph):
            return
        
        should_be_bold = None
        section_type = None
        
        # 1. ชื่อบท - ควรเป็น Bold
        if self._is_main_heading(text, paragraph):
            should_be_bold = True
            section_type = "ชื่อบท"
        # 2. หัวข้อใหญ่ (1. 2. 3.) - ควรเป็น Bold
        elif self._is_major_heading(text, paragraph):
            should_be_bold = True
            section_type = "หัวข้อใหญ่"
        # 3. หัวข้อย่อยทุกระดับ - ควรเป็นตัวปกติ (ไม่หนา)
        elif self._is_any_sub_heading(text):
            should_be_bold = False
            section_type = "หัวข้อย่อย"
        
        # ตรวจสอบความหนาเฉพาะเมื่อมีการกำหนด
        if should_be_bold is not None and section_type:
            # ตรวจสอบความหนาใน run ที่มีเนื้อหาจริง (ไม่ใช่ error message)
            content_runs_bold_status = []
            
            for run in paragraph.runs:
                run_text = run.text.strip()
                if run_text and "[❌ ข้อผิดพลาด:" not in run_text:
                    content_runs_bold_status.append(bool(run.font.bold))
            
            # ถ้ามี content runs อย่างน้อย 1 run
            if content_runs_bold_status:
                # ตรวจสอบว่า run หลักควรเป็น bold หรือไม่
                most_common_bold_status = max(set(content_runs_bold_status), key=content_runs_bold_status.count)
                
                if should_be_bold and not most_common_bold_status:
                    self.errors.append({
                        'paragraph': para_num,
                        'type': 'ความหนาฟอนต์ไม่ถูกต้อง',
                        'description': f'{section_type}ควรเป็นตัวหนา (Bold)',
                        'paragraph_obj': paragraph
                    })
                    self._highlight_paragraph(paragraph)
                    self._add_comment_to_paragraph(paragraph, f"{section_type}ควรเป็นตัวหนา (Bold)")
                elif not should_be_bold and most_common_bold_status:
                    self.errors.append({
                        'paragraph': para_num,
                        'type': 'ความหนาฟอนต์ไม่ถูกต้อง',
                        'description': f'{section_type}ควรเป็นตัวปกติ (ไม่หนา)',
                        'paragraph_obj': paragraph
                    })
                    self._highlight_paragraph(paragraph)
                    self._add_comment_to_paragraph(paragraph, f"{section_type}ควรเป็นตัวปกติ (ไม่หนา)")

    def _check_references(self, doc):
        """ตรวจสอบการอ้างอิงตารางและรูปภาพ"""
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # ตรวจสอบว่าเป็น caption หรือไม่จาก style
            is_caption = self._is_caption_style(paragraph)
            
            if is_caption:
                # ถ้าเป็น caption แล้วข้ามการตรวจสอบการเยื้อง
                continue
            
            # ตรวจสอบการอ้างอิงตารางในเนื้อหา (ไม่ใช่ caption)
            if self._contains_table_reference(text) and not is_caption:
                if not self._is_valid_table_reference_format(text):
                    self.errors.append({
                        'paragraph': i+1,
                        'type': 'รูปแบบการอ้างอิงตารางไม่ถูกต้อง',
                        'description': 'การอ้างอิงตารางควรเป็น "ตารางที่ X" หรือ "ตารางที่ X.X"',
                        'paragraph_obj': paragraph
                    })
                    self._highlight_paragraph(paragraph)
                    self._add_comment_to_paragraph(paragraph, 'การอ้างอิงตารางควรเป็น "ตารางที่ X" หรือ "ตารางที่ X.X"')
            
            # ตรวจสอบการอ้างอิงรูปภาพในเนื้อหา (ไม่ใช่ caption)
            if self._contains_picture_reference(text) and not is_caption:
                if not self._is_valid_picture_reference_format(text):
                    self.errors.append({
                        'paragraph': i+1,
                        'type': 'รูปแบบการอ้างอิงรูปภาพไม่ถูกต้อง',
                        'description': 'การอ้างอิงรูปภาพควรเป็น "ภาพที่ X" หรือ "ภาพที่ X.X"',
                        'paragraph_obj': paragraph
                    })
                    self._highlight_paragraph(paragraph)
                    self._add_comment_to_paragraph(paragraph, 'การอ้างอิงรูปภาพควรเป็น "ภาพที่ X" หรือ "ภาพที่ X.X"')

    # ฟังก์ชันช่วยเหลือ
    def _is_caption_style(self, paragraph):
        """ตรวจสอบว่าเป็น caption style หรือไม่"""
        try:
            if paragraph.style and hasattr(paragraph.style, 'name'):
                style_name = paragraph.style.name.lower()
                if 'caption' in style_name:
                    return True
        except:
            pass
        return False

    def _contains_table_reference(self, text):
        """ตรวจสอบว่าข้อความมีการอ้างอิงตารางหรือไม่"""
        patterns = [
            r'ตารางที่\s*\d+',
            r'ตารางที่\s*\d+\.\d+',
            r'table\s*\d+',
            r'Table\s*\d+',
        ]
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _contains_picture_reference(self, text):
        """ตรวจสอบว่าข้อความมีการอ้างอิงรูปภาพหรือไม่"""
        patterns = [
            r'ภาพที่\s*\d+',
            r'ภาพที่\s*\d+\.\d+',
            r'รูปที่\s*\d+',
            r'รูปที่\s*\d+\.\d+',
            r'figure\s*\d+',
            r'Figure\s*\d+',
        ]
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _is_valid_table_reference_format(self, text):
        """ตรวจสอบรูปแบบการอ้างอิงตารางว่าถูกต้องหรือไม่"""
        valid_patterns = [
            r'^ตารางที่\s*\d+\s',
            r'^ตารางที่\s*\d+\.\d+\s',
            r'^ตารางที่\s*\d+',
            r'^ตารางที่\s*\d+\.\d+',
        ]
        for pattern in valid_patterns:
            if re.search(pattern, text):
                return True
        return False
    
    def _is_valid_picture_reference_format(self, text):
        """ตรวจสอบรูปแบบการอ้างอิงรูปภาพว่าถูกต้องหรือไม่"""
        valid_patterns = [
            r'^ภาพที่\s*\d+\s',
            r'^ภาพที่\s*\d+\.\d+\s',
            r'^ภาพที่\s*\d+',
            r'^ภาพที่\s*\d+\.\d+',
        ]
        for pattern in valid_patterns:
            if re.search(pattern, text):
                return True
        return False

    def check_document_object(self, doc):
        """ตรวจสอบ Document object ที่ผ่านเข้ามา"""
        self.errors = []
        self.tables = []
        self.pictures = []
        
        # ตรวจหาตารางและรูปภาพ
        self._extract_tables_and_pictures(doc)
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                # ข้ามหน้าปกในการตรวจสอบปกติ
                if not self._is_cover_page_section(paragraph.text, i):
                    self._check_paragraph(paragraph, i + 1, i)
        
        # ตรวจสอบการอ้างอิงตารางและรูปภาพ
        self._check_references(doc)
        
        return self.errors
    
    def _get_paragraph_font_size(self, paragraph):
        """ดึงขนาดฟอนต์จากย่อหน้า"""
        for run in paragraph.runs:
            if run.text.strip() and "[❌ ข้อผิดพลาด:" not in run.text:
                if run.font.size and hasattr(run.font.size, 'pt'):
                    return run.font.size.pt
        
        try:
            if paragraph.style and paragraph.style.font and paragraph.style.font.size:
                if hasattr(paragraph.style.font.size, 'pt'):
                    return paragraph.style.font.size.pt
        except:
            pass
            
        return None
    
    def _is_paragraph_bold(self, paragraph):
        """ตรวจสอบว่าย่อหน้าเป็นตัวหนาหรือไม่"""
        for run in paragraph.runs:
            if run.text.strip() and "[❌ ข้อผิดพลาด:" not in run.text:
                if run.font.bold:
                    return True
        return False
    
    def _check_paragraph(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบย่อหน้าแต่ละย่อหน้า"""
        self._check_thai_font(paragraph, para_num, paragraph_index)
        self._check_font_sizes(paragraph, para_num, paragraph_index)
        self._check_font_bold(paragraph, para_num, paragraph_index)
        self._check_alignment(paragraph, para_num, paragraph_index)
        self._check_indentation_detailed(paragraph, para_num, paragraph_index)
        self._check_spacing(paragraph, para_num, paragraph_index)
        self._check_heading_numbering(paragraph, para_num, paragraph_index)
    
    def _is_cover_page_section(self, text, paragraph_index):
        """ตรวจสอบว่าเป็นส่วนของหน้าปกหรือไม่"""
        # เพิ่มขอบเขตการตรวจสอบหน้าปกให้มากขึ้น (รองรับ 2 หน้า)
        if paragraph_index > 50:  # เพิ่มจาก 25 เป็น 50
            return False
        
        # ตรวจสอบคำสำคัญที่บ่งบอกว่าไม่ใช่หน้าปก
        non_cover_keywords = ['บทคัดย่อ', 'abstract', 'สารบัญ', 'บทที่', 'กิตติกรรม', 'acknowledgement', 'table of contents', 'contents', 'ประกาศนียบัตร', 'certificate']
        for keyword in non_cover_keywords:
            if keyword.lower() in text.lower():
                return False
        
        # ตรวจสอบคำสำคัญที่บ่งบอกว่าเป็นหน้าปก
        cover_keywords_from_standards = self.standards.get('cover_sections', [])
        for keyword in cover_keywords_from_standards:
            if keyword in text:
                return True
        
        # ตรวจสอบรหัสนักศึกษา (รูปแบบ ตัวเลข 9-10 หลัก ตามด้วย - และตัวเลข 1 หลัก)
        if re.search(r'\d{9,10}-\d', text):
            return True
        
        # ตรวจสอบคำสำคัญเพิ่มเติมสำหรับหน้าปก
        additional_cover_keywords = [
            'cs/it', 'cs ', 'ภาคเรียนที่', 'ปีการศึกษา', 'เดือน', 'พ.ศ.', 'ครั้งที่', 
            'วิทยาลัยการคอมพิวเตอร์', 'วิทยาลัยคอมพิวเตอร์', 'มหาวิทยาลัยขอนแก่น', 
            'โครงงาน', 'เอกสารโครงงาน', 'รายงานความก้าวหน้า', 'โดย', 'อาจารย์ที่ปรึกษา',
            'รายงานนี้เป็นส่วนหนึ่ง', 'การศึกษาวิชา'
        ]
        
        for keyword in additional_cover_keywords:
            if keyword.lower() in text.lower():
                return True
                
        return False
    
    def _check_spacing(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบระยะห่างบรรทัด"""
        if not paragraph.text.strip():
            return
            
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing and line_spacing != 1.0:
            if isinstance(line_spacing, (int, float)) and line_spacing > 1.1:
                self.errors.append({
                    'paragraph': para_num,
                    'type': 'ระยะห่างบรรทัดไม่ถูกต้อง',
                    'description': f'ควรใช้ระยะห่างบรรทัด 1 เท่า (Single) แต่พบ {line_spacing}',
                    'paragraph_obj': paragraph
                })
                self._highlight_paragraph(paragraph)
                self._add_comment_to_paragraph(paragraph, f"ควรใช้ระยะห่างบรรทัด 1 เท่า (Single) - ปัจจุบัน {line_spacing}")
    
    def _check_heading_numbering(self, paragraph, para_num, paragraph_index):
        """ตรวจสอบการใช้เลขหัวข้อ"""
        text = paragraph.text.strip()
        
        if re.match(r'^\d+(\.\d+)*\s*$', text):
            self.errors.append({
                'paragraph': para_num,
                'type': 'รูปแบบหัวข้อไม่ถูกต้อง',
                'description': 'หัวข้อควรมีเนื้อหาตามหลังเลขหัวข้อ',
                'paragraph_obj': paragraph
            })
            self._highlight_paragraph(paragraph)
            self._add_comment_to_paragraph(paragraph, "หัวข้อควรมีเนื้อหาตามหลังเลขหัวข้อ")
    
    def _should_skip_indent_check(self, text):
        """ตรวจสอบว่าควรข้ามการตรวจสอบการเยื้องหรือไม่"""
        # เพิ่มรายการที่ไม่ต้องตรวจสอบการเยื้อง
        additional_skip_indent = [
            'keywords:', 'keyword:', 'คำสำคัญ:', 'คำสำคัญ :',
            'หน้า', 'ผู้จัดทำ', 'page', 
            'ตารางที่', 'ภาพที่', 'รูปที่', 'table', 'figure'
        ]
        
        text_lower = text.lower()
        
        # ตรวจสอบรายการเดิม
        for section in self.standards['no_indent_sections']:
            if section in text:
                return True
        
        # ตรวจสอบรายการเพิ่มเติม
        for item in additional_skip_indent:
            if item in text_lower:
                return True
                
        return False
    
    def _should_skip_font_check(self, text):
        """ตรวจสอบว่าควรข้ามการตรวจสอบฟอนต์หรือไม่"""
        for section in self.standards['skip_font_check']:
            if section in text:
                return True
        return False
    
    def _should_skip_bold_check(self, text):
        """ตรวจสอบว่าควรข้ามการตรวจสอบความหนาหรือไม่"""
        # เพิ่มรายการที่ไม่ต้องตรวจสอบความหนา
        additional_skip_bold = [
            'keywords:', 'keyword:', 'คำสำคัญ:', 'คำสำคัญ :',
            'ขอขอบพระคุณ', 'ขอขอบคุณ', 'คุณพ่อ', 'คุณแม่',
            'ผู้จัดทำ', 'หน้า'
        ]
        
        text_lower = text.lower()
        
        # ตรวจสอบรายการเดิม
        for section in self.standards['skip_bold_check']:
            if section in text:
                return True
        
        # ตรวจสอบรายการเพิ่มเติม
        for item in additional_skip_bold:
            if item in text_lower:
                return True
        
        # ข้ามเนื้อหาในกิตติกรรมประกาศ
        if self._is_acknowledgement_content(text):
            return True
                
        return False
    
    def _add_comment_to_paragraph(self, paragraph, comment_text):
        """เพิ่ม comment ที่ท้ายย่อหน้า"""
        if "[❌ ข้อผิดพลาด:" in paragraph.text:
            return
        
        new_run = paragraph.add_run(f" [❌ ข้อผิดพลาด: {comment_text}]")
        new_run.font.color.rgb = RGBColor(255, 0, 0)
        new_run.font.bold = True
        new_run.font.size = Pt(12)
    
    def _highlight_paragraph(self, paragraph):
        """ทำ highlight ย่อหน้า"""
        for run in paragraph.runs:
            if run.text.strip() and "[❌ ข้อผิดพลาด:" not in run.text:
                run.font.highlight_color = 7
    
    def get_error_statistics(self):
        """สร้างสถิติข้อผิดพลาด"""
        error_types = {}
        for error in self.errors:
            error_type = error['type']
            error_types[error_type] = error_types.get(error_type, 0) + 1
        return error_types